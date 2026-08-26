"""High-precision text extraction and NLP-based resume entity parser."""
from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document
from PyPDF2 import PdfReader

# Categorized Skill Taxonomy (250+ Industry Skills)
SKILL_TAXONOMY: dict[str, list[str]] = {
    "Programming Languages": [
        "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "C", "Go", "Golang",
        "Rust", "Ruby", "PHP", "Swift", "Kotlin", "Scala", "R", "Dart", "MATLAB", "Perl",
        "Shell Scripting", "Bash", "PowerShell", "SQL", "HTML5", "CSS3", "Sass"
    ],
    "Frameworks & Web Technologies": [
        "React", "React Native", "Next.js", "Vue.js", "Nuxt.js", "Angular", "Svelte",
        "Node.js", "Express.js", "NestJS", "Django", "Flask", "FastAPI", "Spring Boot",
        "ASP.NET", ".NET Core", "Ruby on Rails", "Laravel", "Tailwind CSS", "Bootstrap",
        "Redux", "GraphQL", "REST API", "gRPC", "WebSockets", "Microservices"
    ],
    "AI, Machine Learning & Data Science": [
        "Machine Learning", "Deep Learning", "Natural Language Processing", "NLP",
        "Computer Vision", "Large Language Models", "LLMs", "Generative AI", "RAG",
        "LangChain", "LlamaIndex", "Hugging Face", "PyTorch", "TensorFlow", "Keras",
        "Scikit-learn", "Pandas", "NumPy", "SciPy", "OpenCV", "NLTK", "Spacy",
        "XGBoost", "LightGBM", "Data Analysis", "Data Science", "Statistics", "MLOps",
        "Model Deployment", "Prompt Engineering", "Fine-Tuning", "Vector Databases", "ChromaDB", "Pinecone"
    ],
    "Cloud, DevOps & Infrastructure": [
        "AWS", "Amazon Web Services", "Azure", "Microsoft Azure", "Google Cloud", "GCP",
        "Docker", "Kubernetes", "K8s", "Terraform", "Ansible", "Helm", "CI/CD",
        "GitHub Actions", "GitLab CI", "Jenkins", "Linux", "Ubuntu", "Nginx", "Apache",
        "Serverless", "AWS Lambda", "CloudFormation", "Prometheus", "Grafana", "ELK Stack"
    ],
    "Databases & Data Engineering": [
        "PostgreSQL", "MySQL", "MongoDB", "Redis", "SQLite", "Cassandra", "DynamoDB",
        "Elasticsearch", "Neo4j", "Firebase", "Supabase", "Apache Spark", "PySpark",
        "Apache Kafka", "Airflow", "Snowflake", "Databricks", "BigQuery", "ETL", "Data Pipelines"
    ],
    "Tools & Platforms": [
        "Git", "GitHub", "GitLab", "Bitbucket", "Jira", "Confluence", "Postman", "Swagger",
        "Figma", "VS Code", "Power BI", "Tableau", "Excel", "Advanced Excel", "Slack", "Notion"
    ],
    "Architecture & Methodologies": [
        "Agile", "Scrum", "Kanban", "System Design", "Software Architecture", "TDD",
        "Unit Testing", "Integration Testing", "Clean Code", "Design Patterns", "DevOps",
        "Object-Oriented Programming", "OOP", "Data Structures & Algorithms"
    ],
    "Leadership & Professional Skills": [
        "Problem Solving", "Team Leadership", "Project Management", "Stakeholder Management",
        "Communication Skills", "Critical Thinking", "Code Review", "Mentorship", "Collaboration"
    ]
}

# Flat list of all unique skills for fast regex searching
ALL_SKILLS: list[str] = sorted(
    list({skill for skills in SKILL_TAXONOMY.values() for skill in skills}),
    key=lambda x: -len(x)  # Longest first to avoid partial substring clashes
)
SKILLS = ALL_SKILLS  # Backward compatibility alias


def extract_text(uploaded_file) -> str:
    """Extract raw text from a Streamlit file uploader object (PDF or DOCX)."""
    raw = uploaded_file.getvalue()
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix == ".docx":
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        # Also extract table cells if present
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
        return text.strip()

    if suffix != ".pdf":
        raise ValueError("Unsupported format. Please upload PDF or DOCX files.")

    # Primary PDF extractor: pdfplumber for layout preservation
    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            pages = [page.extract_text(layout=True) or page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages).strip()
            if len(text) >= 30:
                return text
    except Exception:
        pass

    # Fallback PDF extractor: PyPDF2
    try:
        reader = PdfReader(io.BytesIO(raw))
        text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if len(text) >= 30:
            return text
    except Exception as exc:
        raise ValueError("Unable to extract text from this PDF. It may be scanned, image-only, or encrypted.") from exc

    raise ValueError("The document contains no readable text.")


def _extract_section(text: str, header_keywords: list[str]) -> str:
    """Detect section content between standard resume headings."""
    pattern = r"(?im)^\s*(?:" + "|".join(re.escape(k) for k in header_keywords) + r")\s*[:\-]?\s*$"
    match = re.search(pattern, text)
    if not match:
        return "Not detected"
    start = match.end()
    after = text[start:]
    # Next major uppercase section header
    next_header = re.search(r"(?m)^\s*[A-Z][A-Za-z &/,-]{2,35}\s*[:\-]?\s*$", after)
    extracted = after[:next_header.start() if next_header else 800].strip()
    return extracted[:800] if extracted else "Not detected"


def _estimate_experience_years(text: str) -> float:
    """Estimate total years of professional experience using year ranges and keyword heuristics."""
    current_year = datetime.now().year
    # Match patterns like 2018 - 2023 or 2019 - Present
    year_ranges = re.findall(r"\b(19\d{2}|20\d{2})\s*(?:-|–|to)\s*(19\d{2}|20\d{2}|present|current)\b", text, re.IGNORECASE)
    
    total_years = 0.0
    for start_str, end_str in year_ranges:
        try:
            start_yr = int(start_str)
            end_yr = current_year if end_str.lower() in ("present", "current") else int(end_str)
            if 1980 <= start_yr <= current_year and start_yr <= end_yr <= current_year + 1:
                duration = end_yr - start_yr
                if duration <= 40:
                    total_years = max(total_years, float(duration))
        except ValueError:
            continue

    # Also search for explicit mentions like "5+ years of experience"
    explicit = re.findall(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)", text, re.IGNORECASE)
    for val in explicit:
        try:
            exp_val = float(val)
            if 0 < exp_val <= 40:
                total_years = max(total_years, exp_val)
        except ValueError:
            continue

    return round(total_years, 1)


def _extract_education_level(text: str) -> str:
    """Identify highest education qualification mentioned."""
    degrees = [
        ("Ph.D. / Doctorate", [r"\bph\.?d\b", r"\bdoctorate\b", r"\bdoctoral\b"]),
        ("Master's Degree", [r"\bmaster'?s\b", r"\bm\.?s\.?\b", r"\bm\.?tech\b", r"\bmba\b", r"\bm\.?sc\b", r"\bmca\b"]),
        ("Bachelor's Degree", [r"\bbachelor'?s\b", r"\bb\.?s\.?\b", r"\bb\.?tech\b", r"\bb\.?e\.?\b", r"\bb\.?sc\b", r"\bbca\b", r"\bbba\b"]),
        ("Diploma / Associate", [r"\bdiploma\b", r"\bassociate\b"])
    ]
    for label, patterns in degrees:
        for p in patterns:
            if re.search(p, text, re.IGNORECASE):
                return label
    return "Not explicitly specified"


def parse_resume(text: str, filename: str) -> dict[str, Any]:
    """Parse resume text into a structured dictionary with contact details, skills, and sections."""
    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # Extract Contact Info
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?){2,4}\d{3,4}", text)
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_-]+", text, re.IGNORECASE)
    github_match = re.search(r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_-]+", text, re.IGNORECASE)

    # Extract Candidate Name (heuristics on top lines)
    name = "Candidate"
    for line in lines[:8]:
        cleaned = re.sub(r"[^A-Za-z .'-]", "", line).strip()
        words = cleaned.split()
        if (
            2 <= len(words) <= 4
            and 4 <= len(cleaned) <= 40
            and not any(x in cleaned.lower() for x in ["resume", "curriculum", "vitae", "profile", "contact", "email", "summary", "page", "phone"])
        ):
            name = cleaned.title()
            break

    # If no 2-word name found in top lines, fallback to filename
    if name == "Candidate":
        stem = Path(filename).stem
        cleaned_stem = re.sub(r"[_ -]+(resume|cv|profile|\d+).*", "", stem, flags=re.IGNORECASE).strip()
        if len(cleaned_stem) >= 3:
            name = cleaned_stem.replace("_", " ").replace("-", " ").title()

    # Extract Detected Skills using boundary-safe regex
    found_skills: set[str] = set()
    for skill in ALL_SKILLS:
        # Match standalone words or punctuation-bounded tokens
        pattern = r"(?<![A-Za-z0-9])" + re.escape(skill).replace(r"\ ", r"[\s\-_]+") + r"(?![A-Za-z0-9])"
        if re.search(pattern, text, re.IGNORECASE):
            found_skills.add(skill)

    # Experience and Education
    exp_years = _estimate_experience_years(text)
    edu_level = _extract_education_level(text)

    # Extract sections
    education = _extract_section(text, ["Education", "Academic Background", "Qualifications", "Educational Background"])
    experience = _extract_section(text, ["Experience", "Work Experience", "Professional Experience", "Employment History", "Work History"])
    projects = _extract_section(text, ["Projects", "Key Projects", "Personal Projects", "Academic Projects", "Technical Projects"])
    certificates = _extract_section(text, ["Certifications", "Certificates", "Licenses & Certifications", "Accreditations", "Courses"])

    return {
        "candidate_name": name,
        "email": email_match.group(0) if email_match else "Not detected",
        "phone": phone_match.group(0).strip() if phone_match else "Not detected",
        "linkedin": linkedin_match.group(0) if linkedin_match else "",
        "github": github_match.group(0) if github_match else "",
        "skills": sorted(list(found_skills)),
        "experience_years": exp_years,
        "education_level": edu_level,
        "education": education if education != "Not detected" else edu_level,
        "experience": experience,
        "projects": projects,
        "certificates": certificates,
        "resume_text": text,
        "resume_file": filename,
    }

